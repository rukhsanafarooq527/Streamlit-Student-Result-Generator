# Copyright (C) 2025 Rukhsana Shaheen  
# All Rights Reserved.  
# This code cannot be copied or distributed without permission.  
# Developed by Rukhsana Shaheen for Growth Mindset Challenge  

import streamlit as st
import pandas as pd
from PIL import Image
import io
from fpdf import FPDF
from docx import Document


# --- Page Title and Configuration ---
st.set_page_config(page_title="Student Result Generator", layout="wide")


# st.markdown(
#     "<h1 style='text-align: center; font-size: 50px; color: black; font-weight: bold;'>Student Result Generator</h1>",
#     unsafe_allow_html=True
# )


#css
st.markdown("""
    <style>
        /* Pure Page aur Text ko Black Karna */
        body, h1, h2, h4, h6, p,button{
            color: black !important;
        }

        /* Sidebar Inputs ko Black Karna */
        [class^="stTextInput"] label, 
        [class^="stNumberInput"] label {
            color: black !important;
        }
          
    </style>
""", unsafe_allow_html=True)

# --- Sidebar / Left Panel ---
st.sidebar.header("Enter Student Details")

institution_name = st.sidebar.text_input("Institution Name")
student_name = st.sidebar.text_input("Student Name")
roll_number = st.sidebar.text_input("Roll Number")
class_name = st.sidebar.text_input("Class")
year = st.sidebar.text_input("Year")

uploaded_file = st.sidebar.file_uploader("Upload Student Photo", type=["jpg", "jpeg", "png"])
student_image = Image.open(uploaded_file) if uploaded_file else None


#MARK OBTAINED 1
num_subjects = st.sidebar.number_input("Number of Subjects", min_value=1, value=3)

subjects, marks_obtained, total_marks = [], [], []
for i in range(num_subjects):
    col1, col2, col3 = st.sidebar.columns(3)
    subjects.append(col1.text_input(f"Subject {i+1}", key=f"sub{i}"))
    marks_obtained.append(col2.number_input(f"Marks {i+1}", min_value=0, max_value=100, value=0, key=f"marks{i}"))
    total_marks.append(col3.number_input(f"Total {i+1}", min_value=1, value=100, key=f"total{i}"))

#MARK OBTAINED 2

# num_subjects = st.sidebar.number_input("Number of Subjects", min_value=1, value=3)

# subjects, marks_obtained, total_marks = [], [], []
# for i in range(num_subjects):
#     col1, col2, col3 = st.sidebar.columns(3)
#     subjects.append(col1.text_input(f"Subject {i+1}", key=f"sub{i}"))

#     # Marks: Default 0.0 show hoga, user integer ya decimal likh sakta hai
#     marks_obtained.append(col2.number_input(f"Marks {i+1}", 
#                                             min_value=0.0, 
#                                             max_value=100.0, 
#                                             value=0.0,   # ✅ Integer ki jagah float diya
#                                             step=0.1, 
#                                             format="%.1f", 
#                                             key=f"marks{i}"))

#     # Total: Default 100 show hoga, user change kar sakta hai
#     total_marks.append(col3.number_input(f"Total {i+1}", 
#                                          min_value=1, 
#                                          value=100,   # ✅ Isko int hi rakha hai
#                                          step=1, 
#                                          key=f"total{i}"))

if st.sidebar.button("Generate Result"):
    total_obtained = sum(marks_obtained)
    total_possible = sum(total_marks)
    percentage = (total_obtained / total_possible) * 100 if total_possible else 0

    grade = (
        "A+" if percentage >= 90 else
        "A" if percentage >= 80 else
        "B" if percentage >= 70 else
        "C" if percentage >= 60 else
        "D" if percentage >= 50 else
        "F"
    )

    df = pd.DataFrame({
        "Subject": subjects, 
        "Marks Obtained": marks_obtained, 
        "Total Marks": total_marks
    })
    df.index = range(1, len(df) + 1)

    # --- Main Content ---
    st.markdown(f"<h1 style='text-align: center;'>{institution_name}</h1>", unsafe_allow_html=True)
    st.markdown(f"<h2 style='text-align: center;'>Student Result</h2>", unsafe_allow_html=True)
    
    st.markdown("### Student Information")
    col1, col2 = st.columns([3, 1])
    with col1:
        st.write(f"**Name:** {student_name}")
        st.write(f"**Roll No:** {roll_number}")
        st.write(f"**Class:** {class_name}")
        st.write(f"**Year:** {year}")
    with col2:
        if student_image:
            st.image(student_image, width=150)

    st.subheader("Subject Wise Result")
    st.table(df)
    
    
    st.subheader("Overall Result")
    styled_text = f"""
    <div style="display: flex; justify-content: space-between; align-items: center; width: 100%; font-weight: bold; padding: 10px; border-radius: 10px; background-color: #f5f5f5;">
        <div style="text-align: left;">
            <span style="font-size: 20px;">Total Marks:</span>
            <span style="font-size: 36px;">{total_obtained} / {total_possible}</span>
        </div>
        <div style="text-align: center;">
            <span style="font-size: 20px;">Percentage:</span>
            <span style="font-size: 36px;">{round(percentage, 2)}%</span>
        </div>
        <div style="text-align: right;">
            <span style="font-size: 20px;">Grade:</span>
            <span style="font-size: 36px; color: {'red' if grade == 'F' else 'green'};">{grade}</span>
        </div>
    </div>
    """
    st.markdown(styled_text, unsafe_allow_html=True)
    
    # Remarks Section
    st.markdown("<h2 style='text-align: center; font-size: 25px; font-weight: bold;'>Remarks</h2>", unsafe_allow_html=True)
    remarks = "Congratulations! You Passed." if percentage >= 40 else "Unfortunately, You Failed."
    st.markdown(f"<h3 style='text-align: center;font-size: 40px; color: {'green' if percentage >= 40 else 'red'};'>{remarks}</h3>", unsafe_allow_html=True)

    # Export Functions
    def generate_pdf():
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, institution_name, ln=True, align='C')
        pdf.cell(200, 10, "Student Result", ln=True, align='C')
        pdf.ln(10)
        
        if student_image:
            image_path = "temp_image.jpg"
            student_image.save(image_path)
            pdf.image(image_path, x=145, y=20, w=30)
        
        pdf.set_font("Arial", size=12)
        pdf.cell(0, 8, f"Name: {student_name}", ln=True)
        pdf.cell(0, 8, f"Roll No: {roll_number}", ln=True)
        pdf.cell(0, 8, f"Class: {class_name}", ln=True)
        pdf.cell(0, 8, f"Year: {year}", ln=True)
        pdf.ln(10)
        
        for sub, marks, total in zip(subjects, marks_obtained, total_marks):
            pdf.cell(60, 10, sub, 1, 0, 'C')
            pdf.cell(40, 10, str(marks), 1, 0, 'C')
            pdf.cell(40, 10, str(total), 1, 1, 'C')
        
        pdf.ln(10)
        pdf.cell(0, 10, f"Total Marks: {total_obtained} / {total_possible}", ln=True)
        pdf.cell(0, 10, f"Percentage: {round(percentage, 2)}%", ln=True)
        pdf.cell(0, 10, f"Grade: {grade}", ln=True)
        pdf.cell(0, 10, "Remarks: " + remarks, ln=True)
        return bytes(pdf.output(dest="S"))

    def generate_word():
        doc = Document()
        doc.add_heading(institution_name, level=1)
        doc.add_heading("Student Result", level=2)
        doc.add_paragraph(f"Name: {student_name}")
        doc.add_paragraph(f"Roll No: {roll_number}")
        doc.add_paragraph(f"Class: {class_name}")
        doc.add_paragraph(f"Year: {year}")
        doc.add_paragraph(f"Total Marks: {total_obtained} / {total_possible}")
        doc.add_paragraph(f"Percentage: {round(percentage, 2)}%")
        doc.add_paragraph(f"Grade: {grade}")
        doc.add_paragraph(f"Remarks: {remarks}")
        word_io = io.BytesIO()
        doc.save(word_io)
        word_io.seek(0)
        return word_io.getvalue()


    pdf_data = generate_pdf()
    word_data = generate_word()
    
    col1, col2 = st.columns(2)
    with col1:
        st.download_button("Export as PDF", pdf_data, f"{student_name}_result.pdf", "application/pdf")
    with col2:
        st.download_button("Export as Word", word_data, f"{student_name}_Result.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")




# # Copyright footer
# st.markdown(
#     "<p style='text-align: center; color: grey;'>Copyright © 2025 | Developed by Rukhsana Shaheen Hussain</p>",
#     unsafe_allow_html=True
# )





st.markdown("""
    <style>
        .footer {
            position: fixed;
            bottom: 0;
            width: 120%;
            text-align: center;
            background-color: white;
            padding: 10px;
        }
    </style>
    <div class='footer'>
        <span style= "color:grey,text:small">Copyright © 2025 | Developed by Rukhsana Shaheen Hussain</span>
    </div>
""", unsafe_allow_html=True)









































